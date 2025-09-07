from docx import Document
from .models.resume_info import ResumeInfo

def render_resume_docx(
    resume: ResumeInfo, 
    output_path="final_resume.docx"
):
    doc = Document()

    # Title
    doc.add_heading(resume.name, level=0)

    # Contact Info
    doc.add_paragraph(f"📧 {resume.contact.email}")
    doc.add_paragraph(f"📱 {resume.contact.phone}")

    # Summary
    if resume.summary:
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(resume.summary)

    # Skills
    if resume.skills:
        doc.add_heading("Skills", level=1)
        for cat in resume.skills:
            doc.add_paragraph(", ".join(cat.items), style="List Bullet")

    # Education
    if resume.education:
        doc.add_heading("Education", level=1)
        for e in resume.education:
            entry = f"{e.degree} - {e.institution} ({e.start_date} - {e.end_date})"
            if e.grade:
                entry += f" | Grade: {e.grade}"
            doc.add_paragraph(entry, style="List Bullet")

    # Experience
    if resume.experiences:
        doc.add_heading("Experience", level=1)
        for ex in resume.experiences:
            entry = f"{ex.job_title} - {ex.company}, {ex.location} ({ex.start_date} - {ex.end_date})"
            doc.add_paragraph(entry, style="List Bullet")
            if ex.responsibilities:
                for r in ex.responsibilities:
                    doc.add_paragraph(f"• {r}", style="List Bullet 2")
            if ex.achievements:
                for a in ex.achievements:
                    doc.add_paragraph(f"✔ {a}", style="List Bullet 2")

    # Certifications
    if resume.certifications:
        doc.add_heading("Certifications", level=1)
        for c in resume.certifications:
            doc.add_paragraph(f"{c.name} – {c.issuer} ({c.date})", style="List Bullet")

    # Languages
    if resume.languages:
        doc.add_heading("Languages", level=1)
        for l in resume.languages:
            doc.add_paragraph(f"{l.name}: {l.proficiency}", style="List Bullet")

    # Save file
    doc.save(output_path)
    return output_path
